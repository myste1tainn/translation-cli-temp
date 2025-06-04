from docx import Document
from docx.oxml import CT_Tbl, CT_P, CT_Br
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK
import openpyxl
from google.cloud import translate_v3 as translate

client = translate.TranslationServiceClient()

def translate_text(value, glossary_id, project_id, location="us-central1"):
    if isinstance(value, str) and value.strip():
        parent = f"projects/{project_id}/locations/{location}"
        glossary = f"{parent}/glossaries/{glossary_id}"

        response = client.translate_text(
            request={
                "parent": parent,
                "contents": [value],
                "mime_type": "text/plain",
                "source_language_code": "en",
                "target_language_code": "es",
                "glossary_config": {"glossary": glossary},
            }
        )

        return response.glossary_translations[0].translated_text
    return value

async def translate_docx(input_path, output_path, glossary_id, project_id):
    doc = Document(input_path)

    element_to_obj = {}
    para_idx, table_idx = 0, 0
    for element in doc._element.body:
        if isinstance(element, CT_P):
            element_to_obj[element] = doc.paragraphs[para_idx]
            para_idx += 1
        elif isinstance(element, CT_Tbl):
            element_to_obj[element] = doc.tables[table_idx]
            table_idx += 1

    #Iterate over all elements in the document (paragraphs and tables in order)
    for element, obj in element_to_obj.items():
        if isinstance(obj, type(doc.paragraphs[0])): # If the element is a paragraph
            for run in obj.runs:
                translated_text = translate_text(run.text, glossary_id, project_id)
                # Check if the run contains a page break
                has_page_break = any(isinstance(break_elem, CT_Br) and break_elem.get(qn('w:type')) == 'page' for break_elem in run._element)

                # Replace the text directly
                run.text = translated_text

                # Re-add the page break if it was present
                if has_page_break:
                    run.add_break(WD_BREAK.PAGE)

        elif isinstance(obj, type(doc.tables[0])):
            for i, row in enumerate(obj.rows):
                for j, cell in enumerate(row.cells):
                    for run in cell.paragraphs[0].runs:
                        translated_text = await translate_text(run.text)
                        run.clear()  # Clear the existing text
                        run.add_text(translated_text)  # Add the translated text
    doc.save(output_path)
    print(f"Translation complete. Saved as {output_path}")

if __name__ == "__main__":
    print("Start converting\n\n")
    # Load Excel
    file_path = "./data/gpsc/notes-cut.docx"
    output_path = "translated-notes-cut.docx"
    glossary_id = "your-glossary-id"
    project_id = "sanguine-healer-450808-s3"
    translate_docx(file_path, output_path, glossary_id, project_id)

    print("Done converting the file\n\n")
     
