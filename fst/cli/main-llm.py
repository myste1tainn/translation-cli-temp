from docx import Document
import os
from openai import AzureOpenAI
import json
import pandas as pd
import uuid

endpoint = os.getenv("ENDPOINT_URL", "https://llm-platform-gateway.azure-api.net")  
deployment = os.getenv("DEPLOYMENT_NAME", "claude-3-5-sonnet-v2")  #gemini-pro-1.5 #claude-3-5-sonnet-v2 #gpt-4o
subscription_key = os.getenv("AZURE_OPENAI_API_KEY", "Bearer 824fd67a3fa8bd3a2f280cc72b7fbcf39de924d893c6a191572b40044078461c")   

client = AzureOpenAI(  
    azure_endpoint=endpoint,  
    api_key=subscription_key,  
    api_version="2024-05-01-preview",  
)  

usage_token = 0

jargon_list_str = """
{
        "สินทรัพย์": "Assets",
        "สินทรัพย์หมุนเวียน": "Current assets",
        "เงินสดและรายการเทียบเท่าเงินสด": "Cash and cash equivalents",
        "เงินลงทุนชั่วคราว": "Current investments",
        "ลูกหนี้การค้าและลูกหนี้อื่น": "Trade and other receivables",
        "เงินให้กู้ยืมระยะสั้น": "Current loans and advances",
        "สินค้าคงเหลือ": "Inventories",
        "สินทรัพย์หมุนเวียนอื่น": "Other current assets",
        "รวมสินทรัพย์หมุนเวียน": "Total current assets",
        "สินทรัพย์ไม่หมุนเวียน": "Non-current assets",
        "เงินลงทุนเผื่อขาย": "Investments held as available-for-sale",
        "เงินลงทุนในบริษัทร่วม": "Investments in associates",
        "เงินลงทุนในบริษัทย่อย": "Investments in subsidiaries",
        "เงินลงทุนในการร่วมค้า": "Investments in joint ventures",
        "เงินลงทุนระยะยาวอื่น": "Other non-current investments",
        "เงินให้กู้ยืมระยะยาว": "Non-current loans and advances",
        "อสังหาริมทรัพย์เพื่อการลงทุน": "Investment property",
        "สินทรัพย์ไม่หมุนเวียนที่ถือไว้เพื่อขาย": "Non-current assets classified as held for sale",
        "ที่ดิน อาคารและอุปกรณ์": "Property, plant and equipment",
        "สินทรัพย์ไม่มีตัวตน": "Intangible assets",
        "สินทรัพย์ไม่หมุนเวียนอื่น": "Other non-current assets",
        "รวมสินทรัพย์ไม่หมุนเวียน": "Total non-current assets",
        "รวมสินทรัพย์": "Total assets",
        "หนี้สินและส่วนของผู้ถือหุ้น": "Liabilities and equity",
        "หนี้สินหมุนเวียน": "Current liabilities",
        "เงินเบิกเกินบัญชีและเงินกู้ยืมระยะสั้นจากสถาบันการเงิน": "Bank overdrafts and short-term borrowing from financial institutions",
        "เจ้าหนี้การค้าและเจ้าหนี้อื่น": "Trade and other payables",
        "ส่วนของหนี้สินระยะยาวที่ถึงกำหนดชำระภายในหนึ่งปี": "Current portion of long-term liabilities",
        "เงินกู้ยืมระยะสั้น": "Short-term borrowings",
        "ภาษีเงินได้ค้างจ่าย": "Current income tax payable",
        "ประมาณการหนี้สินระยะสั้น": "Short-term provisions",
        "หนี้สินหมุนเวียนอื่น": "Other current liabilities",
        "รวมหนี้สินหมุนเวียน": "Total current liabilities",
        "หนี้สินไม่หมุนเวียน": "Non-current liabilities",
        "เงินกู้ยืมระยะยาว": "Long-term borrowings",
        "ประมาณการหนี้สินผลประโยชน์พนักงาน": "Employee benefit obligation",
        "ประมาณการหนี้สินระยะยาว": "Long-term provisions",
        "หนี้สินไม่หมุนเวียนอื่น": "Other non-current liabilities",
        "รวมหนี้สินไม่หมุนเวียน": "Total non-current liabilities",
        "รวมหนี้สิน": "Liabilities",
        "ส่วนของผู้ถือหุ้น": "Shareholders' equity",
        "ทุนเรือนหุ้น": "Share capital",
        "ทุนจดทะเบียน": "Authorized share capital",
        "หุ้นบุริมสิทธิ": "Preference shares",
        "จำนวนหุ้น": "Number of shares",
        "มูลค่าต่อหุ้น": "Par value",
        "หุ้นสามัญ": "Ordinary shares",
        "ทุนที่ชำระแล้ว": "Paid-up share capital",
        "ส่วนเกินมูลค่าหุ้น": "Share premium account",
        "ส่วนเกินมูลค่าหุ้นบุริมสิทธิ": "Share premium account-preference shares",
        "ส่วนเกินมูลค่าหุ้นสามัญ": "Share premium account-ordinary shares",
        "กำไร (ขาดทุน) สะสม": "Retained earnings",
        "จัดสรรแล้ว": "Appropriated",
        "ทุนสำรองตามกฎหมาย": "Legal reserves",
        "อื่นๆ": "Others",
        "ยังไม่ได้จัดสรร": "Unappropriated",
        "องค์ประกอบอื่นของส่วนของผู้ถือหุ้น": "Other components of shareholders' equity",
        "รวมส่วนของผู้ถือหุ้น": "Total shareholders' equity",
        "รวมหนี้สินและส่วนของผู้ถือหุ้น": "Total liabilities and shareholders' equity",
        "งบกำไรขาดทุน": "Statement of income",
        "รายได้": "Revenue",
        "รายได้จากการขายหรือการให้บริการ": "Revenues from sales or revenues from services",
        "รายได้ดอกเบี้ย": "Interest income",
        "รายได้ค่าก่อสร้าง": "Contract revenue",
        "รายได้อื่น": "Other incomes",
        "รวมรายได้": "Total revenue",
        "ค่าใช้จ่าย": "Expenses",
        "ต้นทุนขายหรือต้นทุนการให้บริการ": "Costs of sale of goods and rendering of services",
        "ต้นทุนการก่อสร้าง": "Contract costs",
        "ค่าใช้จ่ายในการขาย": "Selling expense",
        "ค่าใช้จ่ายในการบริหาร": "Administrative expense",
        "ค่าใช้จ่ายอื่น": "Other expenses",
        "รวมค่าใช้จ่าย": "Total expenses",
        "กำไร (ขาดทุน) ก่อนต้นทุนทางการเงินและค่าใช้จ่ายภาษีเงินได้": "Profit (loss) before finance costs and income tax expense",
        "ต้นทุนทางการเงิน": "Finance costs",
        "กำไร (ขาดทุน) ก่อนค่าใช้จ่ายภาษีเงินได้": "Profit (loss) before income tax expense",
        "ค่าใช้จ่ายภาษีเงินได้": "Income tax expense",
        "กำไร (ขาดทุน) สุทธิ": "Net profit (loss)",
        "งบแสดงการเปลี่ยนแปลงส่วนของผู้ถือหุ้น": "Statement of changes in shareholders' equity",
        "ยอดคงเหลือต้นงวด": "Equity balance at beginning of period",
        "ผลกระทบของการเปลี่ยนแปลงนโยบายการบัญชี": "Financial effects of changes in accounting policies",
        "ผลสะสมจากการแก้ไขข้อผิดพลาดทางการบัญชี": "Financial effects of corrections of accounting errors",
        "ยอดคงเหลือที่ปรับปรุงแล้ว": "Balance-restated",
        "การเปลี่ยนแปลงในส่วนของผู้ถือหุ้น": "Changes in shareholders' equity",
        "การเพิ่ม (ลด) หุ้นสามัญ": "Increase (decrease)-ordinary shares",
        "การเพิ่ม (ลด) หุ้นบุริมสิทธิ": "Increase (decrease)-preference shares",
        "เงินปันผล": "Dividends",
        "ผลต่างของอัตราแลกเปลี่ยนจากการแปลงค่างบการเงิน": "Exchange difference on translating financial statements",
        "ผลกำไร (ขาดทุน) จากการวัดมูลค่าเงินลงทุนเผื่อขาย": "Gains (losses) on remeasuring available-for-sale investments",
        "ผลกำไร (ขาดทุน) ที่ยังไม่เกิดขึ้นจริงอื่น": "Other unrealized gains (losses)",
        "รวมการเปลี่ยนแปลงส่วนของผู้ถือหุ้น": "Total changes in shareholders' equity",
        "ยอดคงเหลือปลายงวด": "Equity balance at ending of period"
}"""

# Generate a unique request ID
request_id = str(uuid.uuid4())

# Define Metadata
metadata = {
    "user_id": "12345",
    "request_id": request_id,
    "file_name": "GULF-Notes",
    "file_format": "docx",
    "purpose": "document_processing",
    "model_used": deployment  # Tracks the OpenAI model used
}

# Define Tags
tags = {
    "category": "document_translation",
    "file_name": "GULF-Notes",
    "file_format": "docx",
    "source": "AI-generated",
    "request_id": request_id,
    "model_used": deployment 
}

async def google_translate_text(value, dest_lang="en"):
    if isinstance(value, str) and value.strip():
        res = await translator.translate(value, dest=dest_lang)
        return res.text
    return value

def contains_thai(text):
    # Check if the text contains any Thai characters
    return any('\u0E00' <= char <= '\u0E7F' for char in text)

def gen_ai_translate_text(value, company):
    if isinstance(value, str) and value.strip() and contains_thai(value):
        return open_ai_translate(value, company)
    return value

def open_ai_translate(value, company):
    response = client.chat.completions.create(
        model=deployment,
        user="claude-gulf-word-2",
        # response_format={"type": "json_object"},
        # extra_body={
        #     "store": True,  # Enables storing completion
        #     "metadata": metadata,  # Attach structured metadata
        #     "tags": tags  # Attach searchable tags
        # },
        messages=[
            { "role": "system", "content": "You are a financial service translator assistant." },
            { "role": "user", "content": [  
                { 
                    "type": "text", 
                    "text": f"""
Your job is to translate input in Thai to output in English. Please be aware that the content you're translating is from financial statement of a company called: {company}.

You should also use the provided jargon word list to ensure accurate translation. If the input matches any word or phrase from the jargon list, use the predefined translation. If no exact match is found, translate the input based on financial terminology.

jargon word list: {jargon_list_str}

input value: {value}

If the term is not found in the jargon word list, translate based on financial context while maintaining accuracy.

Please give the answer in json format. Respond only with valid JSON. Do not write an introduction or summary:
{{
    "translated_value": "value"
}}

Example: 
For the request where the the input value: "ยอดคงเหลือต้นงวด"

The response should be:

{{
 "translated_value": "Equity balance at beginning of period"
}}
                    """ 
                }
            ] } 
        ],
        max_completion_tokens=1000
    )

    global usage_token
    usage_token += response.usage.completion_tokens

    json_response = response.choices[0].message.content
    response = json.loads(json_response)
    result =  response.get("translated_value")
    print(f"translating {value}, translated {result}")
    return result

def convert_using_docx(file_path, output_path, company):
    doc = Document(file_path)

    jargon_list_json = json.loads(jargon_list_str)
    
    # Translate paragraphs
    for paragraph in doc.paragraphs:
        if isinstance(paragraph.text, str) and paragraph.text.strip():
            if jargon_list_json.get(paragraph.text, None) != None:
                translated_text = jargon_list_json.get(paragraph.text)
            else: 
                translated_text = gen_ai_translate_text(paragraph.text, company)
            
            paragraph.clear()
            paragraph.add_run(translated_text)
    
    # Translate tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if isinstance(paragraph.text, str) and paragraph.text.strip():
                        if jargon_list_json.get(paragraph.text, None) != None:
                            translated_text = jargon_list_json.get(paragraph.text)
                        else: 
                            translated_text = gen_ai_translate_text(paragraph.text, company)

                        paragraph.clear()
                        paragraph.add_run(translated_text)
    
    doc.save(output_path)

if __name__ == "__main__":
    print("Start converting\n\n")
    # Load Word Document
    file_path = "./data/gulf/notes-cut.docx"
    output_path = "translated_gulf_notes2_claude-3-5-sonnet-v2.docx" #claude-3-5-sonnet-v2 #gemini-pro-1.5 #gpt-4o

    startTime = pd.Timestamp.now()
    print(f"##### Start time: {startTime}")

    convert_using_docx(file_path, output_path, "Global Power Synergy Public Company Limited")

    endTime = pd.Timestamp.now()
    print(f"##### End time: {endTime}")
    print(f"######### Total time: {endTime - startTime} | Total token: {usage_token}")

    print("Done converting the file\n\n")
